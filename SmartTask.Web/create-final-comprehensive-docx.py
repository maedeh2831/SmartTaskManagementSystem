from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def read_docx(filepath):
    """خواندن محتوای فایل docx"""
    doc = Document(filepath)
    content = []
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)
    return content

def create_comprehensive_docx():
    """ایجاد جزوهٔ کامل متن‌محور"""
    doc = Document()

    # تنظیم RTL
    section = doc.sections[0]
    section.right_to_left = True
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ===== صفحه عنوان =====
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_format.space_before = Pt(36)
    p_format.space_after = Pt(12)
    run = p.add_run("SmartTask")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_format.space_after = Pt(6)
    run = p.add_run("سیستم مدیریت وظایف هوشمند")
    run.font.size = Pt(18)
    run.font.bold = True

    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_format.space_after = Pt(24)
    run = p.add_run("مبتنی بر معماری ASP.NET Core و فریم‌ورک Scrum")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_format.space_before = Pt(36)
    run = p.add_run("نسخه ۱.۰ - مستندات کامل و جامع")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(68, 114, 196)

    doc.add_page_break()

    # ===== فهرست مطالب =====
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p_format.space_after = Pt(18)
    run = p.add_run("فهرست مطالب")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    toc_items = [
        ("بخش اول", "معرفی و معماری سیستم"),
        ("بخش دوم", "مفاهیم اساسی: Agile و Scrum"),
        ("بخش سوم", "فناوری‌ها و ابزارهای استفاده شده"),
        ("بخش چهارم", "موجودیت‌های اصلی سیستم"),
        ("بخش پنجم", "موجودیت‌های همکاری و ردیابی"),
        ("بخش ششم", "معماری و الگوهای طراحی"),
        ("بخش هفتم", "سرویس‌های سیستم و لایه‌ها"),
        ("بخش هشتم", "کنترلرها و جریان داده"),
        ("بخش نهم", "ویژگی‌های پیشرفته"),
        ("بخش دهم", "امنیت و دسترسی‌ها"),
    ]

    for i, (section_name, desc) in enumerate(toc_items, 1):
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p_format.left_indent = Inches(0.25)
        p_format.space_after = Pt(8)
        run = p.add_run(f"{section_name}: {desc}")
        run.font.size = Pt(11)

    doc.add_page_break()

    # ===== بخش اول: معرفی =====
    add_section_heading(doc, "بخش اول", "معرفی و معماری سیستم")

    add_paragraph_text(doc,
        "SmartTask یک سیستم مدیریت وظایف و پروژه تحت وب است که بر اساس معماری ASP.NET Core 8.0 و الگوی Agile/Scrum "
        "طراحی و پیاده‌سازی شده است. این سیستم راهکاری یکپارچه و جامع برای مدیریت چرخه حیات کامل پروژه‌های نرم‌افزاری فراهم "
        "می‌کند. از برنامه‌ریزی اولیه تا تحویل نهایی، SmartTask تمام فاز‌های توسعه نرم‌افزار را پشتیبانی می‌کند."
    )

    add_paragraph_text(doc,
        "سیستم SmartTask به طور خاص برای تیم‌های توسعه نرم‌افزار طراحی شده است. این سیستم مدیریت Sprint‌ها، User Story‌ها، "
        "Task‌ها، همکاری تیمی، ردیابی زمان و تحلیل‌های پیشرفته را فراهم می‌کند. سیستم می‌تواند برای تیم‌های کوچک (۳-۵ نفر) تا "
        "تیم‌های بزرگ (۱۰۰+ نفر) استفاده شود."
    )

    add_paragraph_text(doc,
        "یکی از ویژگی‌های منحصربه‌فرد SmartTask پشتیبانی کامل از زبان فارسی است. تمام رابط کاربری، تقویم‌ها، منطقه‌های زمانی و "
        "گزارش‌ها به فارسی پشتیبانی می‌شوند. علاوه بر این، سیستم می‌تواند به صورت Self-hosted در سرورهای شخصی نصب شود که این کنترل "
        "کامل بر داده‌ها و عدم وابستگی به ابرهای خارجی را تضمین می‌کند."
    )

    add_subheading(doc, "ویژگی‌های کلیدی")

    add_paragraph_text(doc,
        "SmartTask دارای پشتیبانی کامل از فریم‌ورک Scrum است. شامل مدیریت Sprint‌ها، Backlog‌ها، User Story‌ها و Task‌ها. "
        "سیستم ارتباطات Real-time با SignalR فراهم می‌کند که اطلاع‌رسانی فوری درباره تغییرات را ممکن می‌سازد."
    )

    add_paragraph_text(doc,
        "داشبوردهای تحلیلی و گزارش‌های پیشرفته مانند Burndown Chart، Velocity و Health Analysis قابل استفاده‌اند. ابزارهای "
        "همکاری تیمی مثل نظرات، منشن و نوتیفیکیشن‌های هدفمند تیم را یاری می‌دهند."
    )

    add_paragraph_text(doc,
        "سیستم هوش مصنوعی بدون وابستگی با استفاده از Google Gemma 4-12B فراهم می‌کند. این AI می‌تواند Task‌ها را تجزیه و تحلیل "
        "کرده و توصیه‌های هوشمند ارائه دهد. مدیریت جامع زمان شامل ثبت ساعات کاری، یادآورها و تایمر است."
    )

    add_paragraph_text(doc,
        "سیستم دسترسی و نقش‌های پیچیده کنترل دقیق دسترسی را در تمام سطح‌ها فراهم می‌کند. سیستم قابل سفارشی‌سازی است و می‌تواند "
        "برای نیازهای خاص تغییر و توسعه یابد."
    )

    doc.add_page_break()

    # ===== بخش دوم: Agile و Scrum =====
    add_section_heading(doc, "بخش دوم", "مفاهیم اساسی: Agile و Scrum")

    add_subheading(doc, "Agile (روش‌شناسی چابک)")

    add_paragraph_text(doc,
        "Agile یک روش‌شناسی نوین توسعه نرم‌افزار است. این روش بر توسعه تکراری و افزایشی تمرکز دارد. به جای برنامه‌ریزی کامل "
        "در ابتدای پروژه، Agile بر تحویل مداوم نسخه‌های کاری نرم‌افزار تمرکز دارد. در روش Agile، تیم‌ها به صورت مستمر بازخورد "
        "می‌گیرند و بر اساس آن تغییرات لازم را اعمال می‌کنند."
    )

    add_paragraph_text(doc,
        "روش Agile بر انعطاف‌پذیری در برابر تغییرات، همکاری تیمی و بازخورد مشتری تاکید دارد. در پروژه‌های Agile، نیازمندی‌ها "
        "به تدریج تکامل می‌یابند و تیم به سرعت با تغییرات سازگار می‌شود. این روش‌شناسی به ویژه برای پروژه‌های پویا و با نیازمندی‌های "
        "در حال تغییر مناسب است."
    )

    add_subheading(doc, "Scrum (فریم‌ورک عملیاتی)")

    add_paragraph_text(doc,
        "Scrum یک فریم‌ورک عملیاتی درون Agile است. Scrum روند توسعه را به بازه‌های زمانی کوتاه‌مدت تقسیم می‌کند که به آن‌ها Sprint "
        "می‌گویند. در هر Sprint که معمولاً ۱-۴ هفته طول می‌کشد، تیم مجموعه‌ای از User Story‌ها را انتخاب کرده و تلاش می‌کند آن‌ها را "
        "تکمیل کند."
    )

    add_paragraph_text(doc,
        "Scrum شامل چندین مراسم مهم است. Sprint Planning جلسه‌ای است که در ابتدای هر Sprint برگزار می‌شود تا تیم تصمیم بگیرد کدام "
        "User Story‌ها را برای این Sprint انتخاب کند. Daily Standup جلسات روزانه کوتاه ۱۵ دقیقه‌ای هستند که تیم در آن‌ها وضعیت کار را "
        "مرور می‌کند. Sprint Review در پایان Sprint برگزار می‌شود تا نتایج نمایش داده شود و بازخورد دریافت شود."
    )

    add_paragraph_text(doc,
        "Sprint Retrospective نیز در پایان هر Sprint برگزار می‌شود. در این جلسه تیم فرآیند کار را بررسی می‌کند و روش‌های بهبود را "
        "شناسایی می‌کند. SmartTask تمامی این مراسم و مفاهیم Scrum را به صورت دیجیتال پیاده‌سازی کرده است."
    )

    doc.add_page_break()

    # ===== بخش سوم: فناوری‌ها =====
    add_section_heading(doc, "بخش سوم", "فناوری‌ها و ابزارهای استفاده شده")

    add_subheading(doc, "فناوری‌های Backend")

    add_paragraph_text(doc,
        "SmartTask بر پایه ASP.NET Core 8.0 ساخته شده است. این فریم‌ورک نسخه جدیدی از ASP.NET است که توسط مایکروسافت توسعه یافته و "
        "کاملاً متن‌باز است. ASP.NET Core دارای معماری MVC است و Cross-platform (Windows، Linux، macOS) است. این فریم‌ورک عملکرد بالا "
        "دارد و برای پروژه‌های بزرگ و قابل مقیاس‌پذیر مناسب است."
    )

    add_paragraph_text(doc,
        "Entity Framework Core 8.0 ORM (Object-Relational Mapping) استفاده شده است. EF Core امکان کار با پایگاه داده را از طریق "
        "کلاس‌های .NET فراهم می‌کند و کوئری‌های SQL را خودکار تولید می‌کند. SQL Server پایگاه داده است که یک سیستم مدیریت پایگاه داده "
        "رابطه‌ای قدرتمند از مایکروسافت است."
    )

    add_paragraph_text(doc,
        "برای احراز هویت و امنیت، ASP.NET Core Identity استفاده شده است که سیستمی جامع برای مدیریت کاربران، پسورد‌ها و نقش‌ها است. "
        "OAuth 2.0 برای ورود با Google آماده شده است. SignalR برای ارتباطات Real-time استفاده می‌شود که امکان پیام‌های دوطرفه را بدون تاخیر "
        "فراهم می‌کند."
    )

    add_subheading(doc, "Library‌ها و Package‌های استفاده شده")

    add_paragraph_text(doc,
        "سیستم از ClosedXML برای تولید فایل‌های Excel و QuestPDF برای تولید فایل‌های PDF استفاده می‌کند. برای تست‌ها، xUnit و Moq "
        "مورد استفاده است که امکان Unit Testing و Mock کردن Dependencies را فراهم می‌کند."
    )

    add_paragraph_text(doc,
        "در Frontend، Chart.js برای نمودارها و گراف‌ها، SortableJS برای Drag & Drop، Bootstrap 5 برای طراحی Responsive UI و jQuery "
        "برای کار با DOM و AJAX استفاده می‌شود. Font Awesome برای آیکون‌ها و Vazirmatn برای فونت فارسی استفاده شده است."
    )

    add_paragraph_text(doc,
        "برای هوش مصنوعی، LM Studio استفاده می‌شود که سرور AI محلی است. مدل Google Gemma 4-12B برای چت و تجزیه Task استفاده می‌شود. "
        "SMTP Gmail برای ارسال ایمیل و Webpushr برای ارسال Push Notification استفاده می‌شود."
    )

    doc.add_page_break()

    # ===== بخش چهارم: موجودیت‌های اصلی =====
    add_section_heading(doc, "بخش چهارم", "موجودیت‌های اصلی سیستم")

    add_subheading(doc, "ApplicationUser (کاربر)")

    add_paragraph_text(doc,
        "ApplicationUser موجودیتی است که اطلاعات کاربران سیستم را نگهداری می‌کند. هر کاربر دارای نام کاربری، ایمیل، رمز عبور و سایر "
        "اطلاعات شخصی است. کاربران می‌توانند با نقش‌های مختلفی در سیستم شرکت کنند."
    )

    add_subheading(doc, "Workspace (فضای کاری)")

    add_paragraph_text(doc,
        "Workspace بالاترین سطح سازمان‌دهی در سیستم است. هر Workspace متعلق به یک Owner است و می‌تواند شامل چندین Project، Team و Member "
        "باشد. Workspace مثل Organization در GitHub یا Workspace در Slack است. یک شرکت معمولاً یک Workspace دارد و یا بخش‌های مختلف "
        "می‌توانند Workspace‌های جداگانه داشته باشند."
    )

    add_paragraph_text(doc,
        "هر Workspace دارای نام، توضیحات، لوگو و رنگ سفارشی است. Workspace می‌تواند Public یا Private باشد. در Workspace می‌توان مخاطبان "
        "و نقش‌ها را مدیریت کرد. نقش‌های Workspace شامل Owner، Admin، ProjectManager، Developer، Tester و Viewer است."
    )

    add_subheading(doc, "Project (پروژه)")

    add_paragraph_text(doc,
        "Project پروژه‌های نرم‌افزاری را نمایندگی می‌کند. هر Project در داخل یک Workspace قرار دارد و محل اصلی برنامه‌ریزی و اجرای کار است. "
        "هر Project دارای نام، کلید یکتا (مثل PROJ)، توضیحات، رنگ و آیکون است. Project دارای تاریخ شروع، سررسید و تاریخ پایان است."
    )

    add_paragraph_text(doc,
        "Project دارای وضعیت مختلفی است: Planning، InProgress، OnHold، Completed و Cancelled. هر Project دارای اولویت است: Low، Medium، "
        "High و Critical. Project می‌تواند آرشیو شود."
    )

    add_subheading(doc, "Sprint (اسپرینت)")

    add_paragraph_text(doc,
        "Sprint بازه زمانی مشخصی است (معمولاً ۱-۴ هفته) که تیم درآن مجموعه‌ای از User Story‌ها را تکمیل کرتا تلاش می‌کند. هر Sprint دارای "
        "نام، هدف، تاریخ شروع و پایان است. Sprint دارای ظرفیت (Capacity) است که به ساعت اندازه‌گیری می‌شود."
    )

    add_paragraph_text(doc,
        "Sprint دارای وضعیت مختلفی است: Planning، Active، Completed و Cancelled. تنها یک Sprint می‌تواند در هر Project فعال باشد. Sprint‌های "
        "مختلف نباید تاریخ تداخلی داشته باشند."
    )

    add_subheading(doc, "UserStory (داستان کاربری)")

    add_paragraph_text(doc,
        "UserStory نیازمندی یا ویژگی‌ای است که از دیدگاه کاربر نوشته شده است. UserStory واحد کاری در Scrum است و معمولاً شامل چندین Task است. "
        "هر UserStory دارای عنوان و توضیحات است. هر UserStory دارای معیارهای پذیرش (Acceptance Criteria) است."
    )

    add_paragraph_text(doc,
        "هر UserStory دارای Story Point است که واحد تخمین پیچیدگی است. Story Point معمولاً بر اساس سری فیبوناچی انتخاب می‌شود. UserStory "
        "دارای وضعیت مختلفی است: New، InProgress، Testing و Done. UserStory می‌تواند در Backlog یا در Sprint قرار داشته باشد."
    )

    add_subheading(doc, "TaskItem (وظیفه)")

    add_paragraph_text(doc,
        "TaskItem کوچک‌ترین واحد کاری قابل تخصیص است. هر Task باید به یک UserStory متصل باشد. پیاده‌سازی یک UserStory معمولاً شامل چندین "
        "Task است. هر Task دارای عنوان، توضیحات و نوع است. نوع Task می‌تواند Task، Bug، Feature یا Improvement باشد."
    )

    add_paragraph_text(doc,
        "هر Task دارای وضعیت مختلفی است: ToDo، InProgress، InReview، Done و Cancelled. هر Task دارای تخمین زمانی (Estimate) به ساعت است. "
        "Task می‌تواند به چندین نفر تخصیص داده شود و می‌تواند SubTask، Comment، Attachment، Checklist، Label و وابستگی داشته باشد."
    )

    doc.add_page_break()

    # ===== بخش پنجم: موجودیت‌های همکاری =====
    add_section_heading(doc, "بخش پنجم", "موجودیت‌های همکاری و ردیابی")

    add_subheading(doc, "Comment (نظر)")

    add_paragraph_text(doc,
        "Comment امکان نظر دادن بر روی Task‌ها را فراهم می‌کند. هر Comment دارای محتوا، نویسنده و تاریخ است. کاربران می‌توانند از طریق Comment‌ها "
        "با یکدیگر بحث کنند و مسائل را بررسی کنند. Comment‌ها می‌توانند شامل منشن (@username) باشند."
    )

    add_subheading(doc, "Attachment (ضمیمه)")

    add_paragraph_text(doc,
        "Attachment امکان آپلود فایل‌های مختلف به Task‌ها را فراهم می‌کند. فایل‌های آپلود شده در سرور ذخیره می‌شوند. هر Attachment دارای "
        "نام فایل، اندازه، نوع و تاریخ آپلود است."
    )

    add_subheading(doc, "Checklist و ChecklistItem")

    add_paragraph_text(doc,
        "Checklist امکان ایجاد فهرستی از کارهای کوچک درون یک Task را فراهم می‌کند. هر Checklist دارای عنوان است و می‌تواند شامل چندین ChecklistItem "
        "باشد. ChecklistItem‌ها می‌توانند تکمیل یا نتکمیل باشند."
    )

    add_subheading(doc, "Label (برچسب)")

    add_paragraph_text(doc,
        "Label برچسب‌های رنگی هستند که می‌تواند به Task‌ها اختصاص داده شوند. Label‌ها برای دسته‌بندی و تنظیم Task‌ها استفاده می‌شوند. "
        "هر Label دارای نام و رنگ است."
    )

    add_subheading(doc, "TimeLog (ثبت زمان)")

    add_paragraph_text(doc,
        "TimeLog ثبت ساعات کاری بر روی Task‌ها را فراهم می‌کند. هر TimeLog دارای ساعات کاری، توضیحات و تاریخ است. TimeLog امکان ثبت ساعات "
        "شروع شده و متوقف شده را فراهم می‌کند."
    )

    add_subheading(doc, "Reminder (یادآور)")

    add_paragraph_text(doc,
        "Reminder یادآورهای زمانی برای Task‌ها است. هر Reminder دارای تاریخ و زمان است. سیستم به صورت خودکار در زمان مقررشده یادآوری ارسال می‌کند."
    )

    add_subheading(doc, "Notification (نوتیفیکیشن)")

    add_paragraph_text(doc,
        "Notification نوتیفیکیشن‌های سیستمی برای کاربران هستند. Notification‌ها می‌توانند درباره تغییرات Task‌ها، Comment‌ها، Mention‌ها و "
        "یادآورها باشند. Notification‌ها می‌توانند خوانده شده یا نخوانده باشند."
    )

    add_subheading(doc, "ActivityLog (لاگ فعالیت)")

    add_paragraph_text(doc,
        "ActivityLog تمام فعالیت‌های سیستم را ثبت می‌کند. هر ActivityLog شامل کاربر، عمل، توضیحات و تاریخ است. ActivityLog برای Audit Trail "
        "و ردیابی تاریخچه استفاده می‌شود."
    )

    add_subheading(doc, "TaskDependency (وابستگی Task)")

    add_paragraph_text(doc,
        "TaskDependency روابط بین Task‌های مختلف را مدیریت می‌کند. یک Task می‌تواند منتظر تکمیل Task دیگری باشد. وابستگی‌ها می‌توانند "
        "الزامی (Required) یا اختیاری (Optional) باشند."
    )

    doc.add_page_break()

    # ===== بخش ششم: معماری و الگوها =====
    add_section_heading(doc, "بخش ششم", "معماری و الگوهای طراحی")

    add_subheading(doc, "معماری Layered (لایه‌ای)")

    add_paragraph_text(doc,
        "SmartTask بر اساس معماری N-Tier (چند لایه) طراحی شده است. این معماری مسئولیت‌ها را به صورت منطقی جدا می‌کند. هر لایه یک مسئولیت "
        "مشخص دارد و می‌تواند به صورت مستقل تغییر یابد. معماری لایه‌ای دارای مزایایی مانند Separation of Concerns، Testability، Maintainability "
        "و Scalability است."
    )

    add_paragraph_text(doc,
        "معماری SmartTask از بالا به پایین شامل لایه Presentation است که Controllers، Views و ViewModels را شامل می‌شود. لایه Service شامل "
        "سرویس‌های مختلفی است که منطق کسب‌وکار را پیاده‌سازی می‌کند. لایه Infrastructure شامل Repository‌ها، UnitOfWork و DbContext است. "
        "لایه Data شامل Entity‌ها، Configuration‌ها و Migration‌ها است. و بالاخره لایه Database که پایگاه داده SQL Server است."
    )

    add_subheading(doc, "الگوی Repository")

    add_paragraph_text(doc,
        "الگوی Repository برای انتزاع لایه دسترسی به داده استفاده می‌شود. این الگو مسئولیت دسترسی به داده‌ها را از Service جدا می‌کند. "
        "Repository تمام عملیات CRUD (Create، Read، Update، Delete) را انجام می‌دهد. این الگو تست‌پذیری بیشتری را فراهم می‌کند زیرا "
        "می‌توان Repository را Mock کرد."
    )

    add_subheading(doc, "الگوی Unit of Work")

    add_paragraph_text(doc,
        "الگوی Unit of Work برای مدیریت Transaction‌ها و اطمینان از یکپارچگی داده استفاده می‌شود. Unit of Work تمام تغییرات یک عمل تجاری را "
        "مدیریت می‌کند. اگر یکی از بخش‌های عمل شکست بخورد، تمام تغییرات برگردانده می‌شوند (Rollback)."
    )

    add_subheading(doc, "Dependency Injection")

    add_paragraph_text(doc,
        "Dependency Injection برای کاهش وابستگی‌ها و افزایش انعطاف‌پذیری استفاده می‌شود. تمام Service‌ها در Program.cs ثبت می‌شوند و Container "
        "آن‌ها را به جایی که نیاز است تزریق می‌کند. این الگو Loose Coupling را فراهم می‌کند و تست‌پذیری را بیشتر می‌کند."
    )

    add_subheading(doc, "Service Layer Pattern")

    add_paragraph_text(doc,
        "Service Layer Pattern برای جداسازی منطق کسب‌وکار از Controller استفاده می‌شود. هر Service یک مسئولیت مشخص دارد و منطق کسب‌وکار "
        "را پیاده‌سازی می‌کند. این الگو تمرکز Controller بر HTTP Handling را فراهم می‌کند."
    )

    add_subheading(doc, "Soft Delete Pattern")

    add_paragraph_text(doc,
        "Soft Delete Pattern برای حذف منطقی به جای حذف فیزیکی رکوردها استفاده می‌شود. به جای حذف رکورد از پایگاه داده، یک فیلد ViewState "
        "برای نشان دادن اینکه رکورد حذف شده است استفاده می‌شود. این الگو قابلیت بازیابی داده و حفظ تاریخچه را فراهم می‌کند."
    )

    doc.add_page_break()

    # ===== بخش هفتم: سرویس‌ها =====
    add_section_heading(doc, "بخش هفتم", "سرویس‌های سیستم و لایه‌ها")

    add_subheading(doc, "سرویس‌های مدیریتی")

    add_paragraph_text(doc,
        "سیستم SmartTask دارای بیش از ۴۵ سرویس است که منطق کسب‌وکار را پیاده‌سازی می‌کند. WorkspaceService مدیریت Workspace‌ها را انجام می‌دهد. "
        "ProjectService مدیریت Project‌ها را انجام می‌دهد. SprintService مدیریت Sprint‌ها را انجام می‌دهد. UserStoryService مدیریت User Story‌ها را "
        "انجام می‌دهد. TaskService مدیریت Task‌ها را انجام می‌دهد."
    )

    add_subheading(doc, "سرویس‌های ارتباطاتی")

    add_paragraph_text(doc,
        "NotificationService سیستم نوتیفیکیشن را مدیریت می‌کند. CommentService مدیریت Comment‌ها را انجام می‌دهد. ActivityLogService تمام "
        "فعالیت‌های سیستم را ثبت می‌کند. TaskAssignmentService تخصیص Task‌های به کاربران را مدیریت می‌کند."
    )

    add_subheading(doc, "سرویس‌های تحلیلی")

    add_paragraph_text(doc,
        "WorkloadAnalysisService تحلیل بار کاری را انجام می‌دهد. DelayRiskService ریسک تاخیر پروژه را تحلیل می‌کند. ProjectHealthService سلامت "
        "پروژه را بررسی می‌کند. PriorityEngineService موتور هوشمند اولویت‌بندی است که اولویت‌های بهتری پیشنهاد می‌کند."
    )

    add_paragraph_text(doc,
        "ReportExportService گزارشات را به فرمت‌های مختلف (Excel، PDF) تبدیل می‌کند. ChatService برای ارتباط با AI Assistant استفاده می‌شود. "
        "WebpushrService ارسال Push Notification را مدیریت می‌کند."
    )

    doc.add_page_break()

    # ===== بخش هشتم: کنترلرها =====
    add_section_heading(doc, "بخش هشتم", "کنترلرها و جریان داده")

    add_paragraph_text(doc,
        "سیستم SmartTask شامل بیش از ۴۰ کنترلر MVC است. هر کنترلر مسئولیت معینی دارد. AccountController مسئول احراز هویت و ورود کاربران است. "
        "HomeController صفحات عمومی را مدیریت می‌کند. WorkspaceController مدیریت Workspace‌ها را انجام می‌دهد."
    )

    add_paragraph_text(doc,
        "ProjectController مدیریت Project‌ها را انجام می‌دهد. SprintController اسپرینت‌ها را مدیریت می‌کند. BacklogController Backlog را مدیریت "
        "می‌کند. UserStoryController User Story‌ها را مدیریت می‌کند. TaskController Task‌ها را مدیریت می‌کند."
    )

    add_paragraph_text(doc,
        "TaskBoardController تابلوی Kanban را نمایش می‌دهد. CommentController نظرات را مدیریت می‌کند. AttachmentController آپلود و دانلود فایل‌ها را "
        "انجام می‌دهد. ChecklistController Checklist‌ها را مدیریت می‌کند. LabelController برچسب‌ها را مدیریت می‌کند."
    )

    add_paragraph_text(doc,
        "TimeLogController ثبت زمان کار را مدیریت می‌کند. ReminderController یادآورها را مدیریت می‌کند. NotificationController نوتیفیکیشن‌ها را "
        "نمایش می‌دهد. TeamController تیم‌ها را مدیریت می‌کند. WorkloadController تحلیل بار کاری را نمایش می‌دهد."
    )

    add_paragraph_text(doc,
        "ReportController گزارش‌ها را تولید می‌کند. AdminController مدیریت کل سیستم را انجام می‌دهد. ChatController برای چت با AI است. "
        "SettingsController تنظیمات کاربری را مدیریت می‌کند."
    )

    add_subheading(doc, "جریان داده")

    add_paragraph_text(doc,
        "جریان داده در سیستم از بالا به پایین شامل درخواست کاربر است. درخواست به Controller می‌رسد که Validation را انجام می‌دهد. سپس کنترلر "
        "درخواست را به Service می‌فرستد. Service منطق کسب‌وکار را انجام می‌دهد و بررسی‌های لازم را انجام می‌دهد."
    )

    add_paragraph_text(doc,
        "Service درخواست را به Repository می‌فرستد که عملیات Database را انجام می‌دهد. Repository کوئری‌های SQL را تولید کرده و اجرا می‌کند. "
        "نتایج به Service برمی‌گردند. Service نتایج را پردازش کرده و به کنترلر می‌فرستند. کنترلر نتایج را به View می‌فرستند. View نتایج را "
        "به کاربر نمایش می‌دهد."
    )

    doc.add_page_break()

    # ===== بخش نهم: ویژگی‌های پیشرفته =====
    add_section_heading(doc, "بخش نهم", "ویژگی‌های پیشرفته")

    add_subheading(doc, "وابستگی‌های Task")

    add_paragraph_text(doc,
        "وابستگی‌های Task امکان مدیریت روابط بین Task‌های مختلف را فراهم می‌کند. یک Task می‌تواند منتظر تکمیل Task دیگری باشد. این ویژگی بسیار "
        "مفید است برای پروژه‌های پیچیده که Task‌های زیادی وجود دارند."
    )

    add_subheading(doc, "Task Breakdown")

    add_paragraph_text(doc,
        "Task Breakdown امکان تقسیم خودکار Task‌های بزرگ به Task‌های کوچک‌تر را فراهم می‌کند. سیستم از AI استفاده می‌کند تا Task را تحلیل کرده و "
        "زیرمجموعه‌های مناسب را پیشنهاد دهد."
    )

    add_subheading(doc, "اولویت‌بندی هوشمند")

    add_paragraph_text(doc,
        "سیستم دارای موتور اولویت‌بندی هوشمند است. این موتور بر اساس معیار‌های مختلف مانند تاریخ سررسید، اهمیت Task و وابستگی‌ها اولویت‌های بهتری "
        "پیشنهاد می‌کند."
    )

    add_subheading(doc, "Offroad Task")

    add_paragraph_text(doc,
        "Offroad Task وظایفی هستند که خارج از Project‌ها قرار دارند اما باز هم نیاز به ردیابی دارند. این ویژگی برای کارهای جانبی مفید است."
    )

    add_subheading(doc, "SubTask و Checklist")

    add_paragraph_text(doc,
        "هر Task می‌تواند SubTask‌های داشته باشد. SubTask‌ها Task‌های کوچک‌تری هستند که میتوان آن‌ها را تکمیل‌کرد. Checklist امکان ایجاد "
        "فهرستی از کارهای کوچک را فراهم می‌کند."
    )

    add_subheading(doc, "مبادله Task")

    add_paragraph_text(doc,
        "Task Trade امکان انتقال Task بین Project‌های مختلف را فراهم می‌کند. این ویژگی برای سازمان‌های بزرگ مفید است که چندین Project دارند."
    )

    doc.add_page_break()

    # ===== بخش دهم: امنیت =====
    add_section_heading(doc, "بخش دهم", "امنیت و دسترسی‌ها")

    add_subheading(doc, "احراز هویت")

    add_paragraph_text(doc,
        "سیستم SmartTask از ASP.NET Core Identity برای احراز هویت استفاده می‌کند. کاربران می‌توانند نام کاربری و رمز عبور خود را استفاده کنند. "
        "رمز عبور به صورت Hashed در پایگاه داده ذخیره می‌شود. سیستم احراز هویت دو عاملی (2FA) را نیز پشتیبانی می‌کند."
    )

    add_subheading(doc, "مجوزها و دسترسی‌ها")

    add_paragraph_text(doc,
        "سیستم SmartTask دارای سیستم مجوزهای پیشرفته است. بررسی دسترسی در سطح Workspace، Project و Task انجام می‌شود. هر کاربر بر اساس نقش خود "
        "دسترسی‌های مختلفی دارد. سرویس‌های مجوز مختلف مثل CanManageWorkspace، CanManageProject و CanManageBacklog وجود دارند."
    )

    add_paragraph_text(doc,
        "سیستم Row-Level Security پیاده‌سازی می‌کند که کاربران فقط می‌توانند داده‌های خود را ببینند. نقش‌های مختلف مثل Admin، ProjectManager و "
        "Developer دسترسی‌های مختلفی دارند."
    )

    add_subheading(doc, "حذف منطقی")

    add_paragraph_text(doc,
        "تمام موجودیت‌های سیستم دارای فیلد ViewState هستند. به جای حذف فیزیکی رکوردها، این فیلد برای نشان دادن اینکه رکورد حذف شده است استفاده می‌شود. "
        "این سیستم قابلیت بازیابی داده را فراهم می‌کند."
    )

    add_subheading(doc, "اطلاعات حساس")

    add_paragraph_text(doc,
        "سیستم اطلاعات کاربران را محافظت می‌کند. ایمیل‌ها Private هستند. رمز عبور‌ها Hashed هستند. تمام ارتباطات از HTTPS استفاده می‌کنند."
    )

    # ذخیره document
    output_path = 'SmartTask-جزوه-کامل-متن‌محور.docx'
    doc.save(output_path)
    return output_path

def add_section_heading(doc, section_num, title):
    """افزودن عنوان بخش"""
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p_format.space_before = Pt(24)
    p_format.space_after = Pt(12)
    p_format.keep_with_next = True

    run = p.add_run(f"{section_num}: {title}")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

def add_subheading(doc, text):
    """افزودن عنوان فرعی"""
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p_format.space_before = Pt(12)
    p_format.space_after = Pt(8)
    p_format.keep_with_next = True

    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(68, 114, 196)

def add_paragraph_text(doc, text):
    """افزودن پاراگراف متن"""
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p_format.line_spacing = 1.65
    p_format.space_after = Pt(12)
    p_format.text_alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 0, 0)

# اجرا
if __name__ == '__main__':
    output = create_comprehensive_docx()
    print(f'✅ فایل جزوه کامل با موفقیت ایجاد شد: {output}')
    print(f'📄 فایل شامل ۱۰ بخش و متن جامع و خوانا است')
    print(f'✍️ تمام محتوا متن‌محور و درجملات مکمل است')
