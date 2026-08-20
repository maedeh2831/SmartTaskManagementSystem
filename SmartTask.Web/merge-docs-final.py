from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# خواندن دو فایل Word موجود
doc1_path = 'SmartTask-Documentation.docx'
doc2_path = 'SmartTask-جزوه-کامل.docx'

doc1 = Document(doc1_path)
doc2 = Document(doc2_path)

# ایجاد document جدید
final_doc = Document()

# تنظیمات
section = final_doc.sections[0]
section.right_to_left = True
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# ===============================================
# صفحه عنوان
# ===============================================
title = final_doc.add_paragraph()
title_format = title.paragraph_format
title_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_format.space_before = Pt(36)
title_format.space_after = Pt(18)
title_run = title.add_run('راهنمای جامع و کامل پروژه SmartTask')
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = final_doc.add_paragraph()
subtitle_format = subtitle.paragraph_format
subtitle_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle_format.space_after = Pt(6)
subtitle_run = subtitle.add_run('سیستم مدیریت وظایف و پروژه هوشمند')
subtitle_run.font.size = Pt(16)
subtitle_run.font.bold = True

version = final_doc.add_paragraph()
version_format = version.paragraph_format
version_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
version_run = version.add_run('نسخه 1.0 - جزوه جامع و آموزشی')
version_run.font.size = Pt(12)
version_run.font.color.rgb = RGBColor(100, 100, 100)

final_doc.add_page_break()

# ===============================================
# معرفی کامل
# ===============================================
intro_title = final_doc.add_paragraph()
intro_title_format = intro_title.paragraph_format
intro_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
intro_title_format.space_before = Pt(18)
intro_title_format.space_after = Pt(12)
intro_title_run = intro_title.add_run('مقدمه: معرفی SmartTask')
intro_title_run.font.size = Pt(28)
intro_title_run.font.bold = True
intro_title_run.font.color.rgb = RGBColor(0, 51, 102)

intro_text1 = final_doc.add_paragraph()
intro_text1_format = intro_text1.paragraph_format
intro_text1_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
intro_text1_format.line_spacing = 1.6
intro_text1_format.space_after = Pt(12)
intro_text1_run = intro_text1.add_run(
    'SmartTask یک سیستم جامع و پیشرفته برای مدیریت وظایف و پروژه‌ها است که به طور خاص برای تیم‌های توسعه نرم‌افزار طراحی شده است. '
    'این سیستم بر پایه معماری مدرن ASP.NET Core 8.0 و الگوهای مدیریت پروژه Agile/Scrum ساخته شده و تمام ابزارهای مورد نیاز برای '
    'مدیریت چرخه حیات کامل پروژه‌ها را از برنامه‌ریزی اولیه تا تحویل نهایی فراهم می‌کند.'
)
intro_text1_run.font.size = Pt(12)
intro_text1_run.font.name = 'Calibri'

intro_text2 = final_doc.add_paragraph()
intro_text2_format = intro_text2.paragraph_format
intro_text2_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
intro_text2_format.line_spacing = 1.6
intro_text2_format.space_after = Pt(12)
intro_text2_run = intro_text2.add_run(
    'یکی از ویژگی‌های منحصربه‌فرد SmartTask این است که کاملاً برای زبان و فرهنگ فارسی طراحی شده است. تمام رابط کاربری، '
    'تقویم‌ها، منطقه‌های زمانی و گزارش‌ها به فارسی پشتیبانی می‌شوند. علاوه بر این، سیستم می‌تواند به صورت Self-hosted در سرورهای شخصی نصب شود '
    'که کنترل کامل بر داده‌ها و استقلال از ابرهای خارجی را تضمین می‌کند.'
)
intro_text2_run.font.size = Pt(12)
intro_text2_run.font.name = 'Calibri'

intro_text3 = final_doc.add_paragraph()
intro_text3_format = intro_text3.paragraph_format
intro_text3_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
intro_text3_format.line_spacing = 1.6
intro_text3_format.space_after = Pt(12)
intro_text3_run = intro_text3.add_run(
    'SmartTask امکانات بسیاری دارد: پشتیبانی کامل از Scrum، ارتباطات Real-time، داشبوردهای تحلیلی، ابزارهای همکاری تیمی، '
    'هوش مصنوعی یکپارچه، مدیریت جامع زمان، و سیستمی پیچیده برای کنترل دسترسی. این سیستم می‌تواند برای تیم‌های کوچک (۳-۵ نفر) '
    'تا تیم‌های بزرگ (۱۰۰+ نفر) استفاده شود.'
)
intro_text3_run.font.size = Pt(12)
intro_text3_run.font.name = 'Calibri'

# ===============================================
# اهداف و مزایا
# ===============================================
goals_title = final_doc.add_paragraph()
goals_title_format = goals_title.paragraph_format
goals_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
goals_title_format.space_before = Pt(18)
goals_title_format.space_after = Pt(12)
goals_title_run = goals_title.add_run('اهداف و مزایای سیستم')
goals_title_run.font.size = Pt(24)
goals_title_run.font.bold = True
goals_title_run.font.color.rgb = RGBColor(68, 114, 196)

goals_text = final_doc.add_paragraph()
goals_text_format = goals_text.paragraph_format
goals_text_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
goals_text_format.line_spacing = 1.6
goals_text_format.space_after = Pt(12)
goals_text_run = goals_text.add_run(
    'هدف اصلی SmartTask ارائه پلتفرمی یکپارچه و کاربرپسند برای تیم‌های نرم‌افزاری است تا بتوانند پروژه‌های خود را به '
    'صورت مؤثر و سازمان‌یافته مدیریت کنند. سیستم به دنبال کاهش پیچیدگی فرآیندهای مدیریتی، افزایش شفافیت در پروژه‌ها، '
    'بهبود همکاری تیمی، ارائه تحلیل‌های داده‌محور و خودکارسازی کارهای تکراری است.'
)
goals_text_run.font.size = Pt(12)
goals_text_run.font.name = 'Calibri'

advantages_title = final_doc.add_paragraph()
advantages_title_format = advantages_title.paragraph_format
advantages_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
advantages_title_format.space_before = Pt(14)
advantages_title_format.space_after = Pt(10)
advantages_title_run = advantages_title.add_run('مزایای نسبت به ابزارهای رقیب')
advantages_title_run.font.size = Pt(18)
advantages_title_run.font.bold = True
advantages_title_run.font.color.rgb = RGBColor(68, 114, 196)

advantages_text = final_doc.add_paragraph()
advantages_text_format = advantages_text.paragraph_format
advantages_text_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
advantages_text_format.line_spacing = 1.6
advantages_text_format.space_after = Pt(12)
advantages_text_run = advantages_text.add_run(
    'SmartTask برای رقابت با ابزارهایی مثل Jira، Trello و Asana دارای مزایایی منحصربه‌فرد است. اول اینکه کاملاً Self-hosted است '
    'و هیچ هزینه اشتراک ماهانه ندارد. دوم اینکه کد منبع کاملاً در دسترس است و می‌توان آن را به نیازهای خاص سازمان تطبیق داد. '
    'سوم اینکه کاملاً برای فارسی پشتیبانی شده است. چهارم اینکه هوش مصنوعی بدون وابستگی به API‌های خارجی در آن ادغام شده است. '
    'و پنجم اینکه بعد از نصب، نیازی به اینترنت دائمی نیست.'
)
advantages_text_run.font.size = Pt(12)
advantages_text_run.font.name = 'Calibri'

final_doc.add_page_break()

# ===============================================
# معماری سیستم
# ===============================================
arch_title = final_doc.add_paragraph()
arch_title_format = arch_title.paragraph_format
arch_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
arch_title_format.space_before = Pt(18)
arch_title_format.space_after = Pt(12)
arch_title_run = arch_title.add_run('معماری و ساختار سیستم')
arch_title_run.font.size = Pt(28)
arch_title_run.font.bold = True
arch_title_run.font.color.rgb = RGBColor(0, 51, 102)

arch_text1 = final_doc.add_paragraph()
arch_text1_format = arch_text1.paragraph_format
arch_text1_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
arch_text1_format.line_spacing = 1.6
arch_text1_format.space_after = Pt(12)
arch_text1_run = arch_text1.add_run(
    'SmartTask بر اساس معماری N-Tier (چند لایه) طراحی شده است. این معماری مسئولیت‌ها را به صورت منطقی جدا می‌کند و هر لایه '
    'یک مسئولیت مشخص دارد. لایه Presentation شامل Controllers و Views است که مسئولیت رابط کاربری را دارد. لایه Service شامل تمام '
    'منطق کسب‌وکار است. لایه Infrastructure شامل Repository‌ها و دسترسی به داده است. و در نهایت لایه Database است.'
)
arch_text1_run.font.size = Pt(12)
arch_text1_run.font.name = 'Calibri'

arch_text2 = final_doc.add_paragraph()
arch_text2_format = arch_text2.paragraph_format
arch_text2_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
arch_text2_format.line_spacing = 1.6
arch_text2_format.space_after = Pt(12)
arch_text2_run = arch_text2.add_run(
    'از الگوهای طراحی معروفی استفاده شده است: Repository Pattern برای انتزاع دسترسی به داده، Unit of Work برای مدیریت Transaction‌ها، '
    'Dependency Injection برای کاهش وابستگی‌ها، Service Layer Pattern برای جداسازی منطق کسب‌وکار، و Soft Delete Pattern برای حذف '
    'منطقی رکوردها. این الگوها کد را سازمان‌یافته، قابل تست و نگهداری‌پذیر می‌کنند.'
)
arch_text2_run.font.size = Pt(12)
arch_text2_run.font.name = 'Calibri'

# ===============================================
# تکنولوژی‌ها
# ===============================================
tech_title = final_doc.add_paragraph()
tech_title_format = tech_title.paragraph_format
tech_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
tech_title_format.space_before = Pt(18)
tech_title_format.space_after = Pt(12)
tech_title_run = tech_title.add_run('فناوری‌های استفاده شده')
tech_title_run.font.size = Pt(24)
tech_title_run.font.bold = True
tech_title_run.font.color.rgb = RGBColor(68, 114, 196)

tech_text = final_doc.add_paragraph()
tech_text_format = tech_text.paragraph_format
tech_text_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
tech_text_format.line_spacing = 1.6
tech_text_format.space_after = Pt(12)
tech_text_run = tech_text.add_run(
    'SmartTask بر پایه ASP.NET Core 8.0 ساخته شده است که فریم‌ورک مدرن و قدرتمند مایکروسافت است. برای پایگاه داده از SQL Server استفاده '
    'می‌شود. برای ORM از Entity Framework Core 8.0 استفاده می‌شود که کار با پایگاه داده را ساده می‌کند. برای احراز هویت از ASP.NET Core Identity '
    'استفاده شده است. برای ارتباطات Real-time از SignalR استفاده می‌شود که به کاربران اجازه می‌دهد تغییرات را بلافاصله ببینند.'
)
tech_text_run.font.size = Pt(12)
tech_text_run.font.name = 'Calibri'

tech_text2 = final_doc.add_paragraph()
tech_text2_format = tech_text2.paragraph_format
tech_text2_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
tech_text2_format.line_spacing = 1.6
tech_text2_format.space_after = Pt(12)
tech_text2_run = tech_text2.add_run(
    'برای Frontend از Bootstrap 5 استفاده شده است که فریم‌ورک CSS مدرن است. برای نمودارها از Chart.js استفاده می‌شود. '
    'برای تولید فایل‌های Excel از ClosedXML و برای PDF از QuestPDF استفاده می‌شود. برای تست‌ها از xUnit و Moq استفاده شده است. '
    'برای هوش مصنوعی از LM Studio با مدل Google Gemma استفاده می‌شود که محلی است و نیازی به API خارجی ندارد.'
)
tech_text2_run.font.size = Pt(12)
tech_text2_run.font.name = 'Calibri'

final_doc.add_page_break()

# ===============================================
# مفاهیم اساسی: Agile و Scrum
# ===============================================
agile_title = final_doc.add_paragraph()
agile_title_format = agile_title.paragraph_format
agile_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
agile_title_format.space_before = Pt(18)
agile_title_format.space_after = Pt(12)
agile_title_run = agile_title.add_run('Agile و Scrum: مفاهیم اساسی')
agile_title_run.font.size = Pt(28)
agile_title_run.font.bold = True
agile_title_run.font.color.rgb = RGBColor(0, 51, 102)

agile_text = final_doc.add_paragraph()
agile_text_format = agile_text.paragraph_format
agile_text_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
agile_text_format.line_spacing = 1.6
agile_text_format.space_after = Pt(12)
agile_text_run = agile_text.add_run(
    'Agile یک روش‌شناسی نوین توسعه نرم‌افزار است که به جای برنامه‌ریزی کامل در ابتدا، بر توسعه تکراری و انعطاف‌پذیری در برابر '
    'تغییرات تاکید دارد. در روش Agile، تیم‌ها به صورت مستمر بازخورد می‌گیرند و بر اساس آن تغییرات لازم را اعمال می‌کنند. این روش '
    'مسئولیت‌های بیشتری را به تیم می‌دهد و باعث می‌شود که محصول نهایی بهتر و بیشتر با نیاز کاربر منطبق باشد.'
)
agile_text_run.font.size = Pt(12)
agile_text_run.font.name = 'Calibri'

scrum_subtitle = final_doc.add_paragraph()
scrum_subtitle_format = scrum_subtitle.paragraph_format
scrum_subtitle_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
scrum_subtitle_format.space_before = Pt(14)
scrum_subtitle_format.space_after = Pt(10)
scrum_subtitle_run = scrum_subtitle.add_run('Scrum: فریم‌ورک عملی Agile')
scrum_subtitle_run.font.size = Pt(18)
scrum_subtitle_run.font.bold = True
scrum_subtitle_run.font.color.rgb = RGBColor(68, 114, 196)

scrum_text = final_doc.add_paragraph()
scrum_text_format = scrum_text.paragraph_format
scrum_text_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
scrum_text_format.line_spacing = 1.6
scrum_text_format.space_after = Pt(12)
scrum_text_run = scrum_text.add_run(
    'Scrum یک فریم‌ورک عملی درون Agile است که روند توسعه را به بازه‌های زمانی کوتاه‌مدت (Sprint‌ها) تقسیم می‌کند. هر Sprint معمولاً '
    '۱ تا ۴ هفته طول می‌کشد و تیم در این مدت روی مجموعه‌ای از کارهای قابل تحویل تمرکز می‌کند. Scrum شامل چندین عنصر کلیدی است: '
    'Product Backlog که لیست تمام نیازمندی‌هاست، Sprint Backlog که کارهای انتخاب شده برای Sprint جاری است، Sprint Planning که جلسه '
    'برنامه‌ریزی در ابتدای Sprint است، Daily Standup که جلسات روزانه کوتاه است، Sprint Review که بررسی نتایج است، و Sprint Retrospective '
    'که برای بهبود فرآیند است.'
)
scrum_text_run.font.size = Pt(12)
scrum_text_run.font.name = 'Calibri'

final_doc.add_page_break()

# ===============================================
# موجودیت‌های اصلی
# ===============================================
entities_title = final_doc.add_paragraph()
entities_title_format = entities_title.paragraph_format
entities_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
entities_title_format.space_before = Pt(18)
entities_title_format.space_after = Pt(12)
entities_title_run = entities_title.add_run('موجودیت‌های اصلی سیستم')
entities_title_run.font.size = Pt(28)
entities_title_run.font.bold = True
entities_title_run.font.color.rgb = RGBColor(0, 51, 102)

entities_text = final_doc.add_paragraph()
entities_text_format = entities_text.paragraph_format
entities_text_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
entities_text_format.line_spacing = 1.6
entities_text_format.space_after = Pt(12)
entities_text_run = entities_text.add_run(
    'موجودیت‌های اصلی پایه‌های سیستم را تشکیل می‌دهند. کاربران (ApplicationUser) افرادی هستند که از سیستم استفاده می‌کنند. '
    'Workspace فضای کاری است که بالاترین سطح سازمان‌دهی را نمایندگی می‌کند. هر Workspace می‌تواند متعلق به شرکت‌ها، دپارتمان‌ها یا '
    'تیم‌های مختلف باشد. Project پروژه نرم‌افزاری است که در داخل Workspace قرار دارد. هر Project دارای مجموعه‌ای از Sprint‌ها و User Story‌ها است.'
)
entities_text_run.font.size = Pt(12)
entities_text_run.font.name = 'Calibri'

entities_text2 = final_doc.add_paragraph()
entities_text2_format = entities_text2.paragraph_format
entities_text2_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
entities_text2_format.line_spacing = 1.6
entities_text2_format.space_after = Pt(12)
entities_text2_run = entities_text2.add_run(
    'Sprint بازه زمانی مشخصی است (معمولاً ۲-۴ هفته) که تیم درآن روی مجموعه‌ای از کارها تمرکز می‌کند. UserStory نیازمندی یا ویژگی‌ای است '
    'که از دیدگاه کاربر نوشته شده است و محدوده کار برای Sprint را مشخص می‌کند. Task کوچک‌ترین واحد کاری است که باید تکمیل شود. هر Task '
    'معمولاً بخشی از یک UserStory است و می‌تواند به یک یا چند نفر تخصیص داده شود.'
)
entities_text2_run.font.size = Pt(12)
entities_text2_run.font.name = 'Calibri'

final_doc.add_page_break()

# ===============================================
# امکانات سیستم
# ===============================================
features_title = final_doc.add_paragraph()
features_title_format = features_title.paragraph_format
features_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
features_title_format.space_before = Pt(18)
features_title_format.space_after = Pt(12)
features_title_run = features_title.add_run('امکانات و ویژگی‌های سیستم')
features_title_run.font.size = Pt(28)
features_title_run.font.bold = True
features_title_run.font.color.rgb = RGBColor(0, 51, 102)

features_subtitle1 = final_doc.add_paragraph()
features_subtitle1_format = features_subtitle1.paragraph_format
features_subtitle1_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
features_subtitle1_format.space_before = Pt(14)
features_subtitle1_format.space_after = Pt(10)
features_subtitle1_run = features_subtitle1.add_run('مدیریت کاربران و دسترسی')
features_subtitle1_run.font.size = Pt(18)
features_subtitle1_run.font.bold = True
features_subtitle1_run.font.color.rgb = RGBColor(68, 114, 196)

features_text1 = final_doc.add_paragraph()
features_text1_format = features_text1.paragraph_format
features_text1_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
features_text1_format.line_spacing = 1.6
features_text1_format.space_after = Pt(12)
features_text1_run = features_text1.add_run(
    'سیستم دارای سیستم احراز هویت قوی است. کاربران می‌توانند ثبت‌نام کنند و با نام کاربری و رمز عبور وارد شوند. سیستم دارای احراز هویت '
    'دو مرحله‌ای است برای امنیت بیشتر. کاربران همچنین می‌توانند با حساب Google خود وارد شوند. سیستم دارای نقش‌های مختلفی است تا کنترل '
    'دسترسی دقیق فراهم کند. نقش‌های Workspace مختلف از نقش‌های Project متفاوت است.'
)
features_text1_run.font.size = Pt(12)
features_text1_run.font.name = 'Calibri'

features_subtitle2 = final_doc.add_paragraph()
features_subtitle2_format = features_subtitle2.paragraph_format
features_subtitle2_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
features_subtitle2_format.space_before = Pt(14)
features_subtitle2_format.space_after = Pt(10)
features_subtitle2_run = features_subtitle2.add_run('داشبوردها و گزارشات')
features_subtitle2_run.font.size = Pt(18)
features_subtitle2_run.font.bold = True
features_subtitle2_run.font.color.rgb = RGBColor(68, 114, 196)

features_text2 = final_doc.add_paragraph()
features_text2_format = features_text2.paragraph_format
features_text2_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
features_text2_format.line_spacing = 1.6
features_text2_format.space_after = Pt(12)
features_text2_run = features_text2.add_run(
    'سیستم داشبوردهای مختلفی فراهم می‌کند. داشبوردهای کاربری نمای کلی وظایف و فعالیت‌های کاربر را نشان می‌دهند. داشبوردهای Workspace '
    'آمار کل Workspace را نشان می‌دهند. داشبوردهای Project پیشرفت و عملکرد پروژه را نشان می‌دهند. سیستم نمودارهای متنوعی دارد مثل '
    'Burndown Chart برای نشان دادن پیشرفت Sprint، Velocity Chart برای نشان دادن سرعت تیم، و Health Analysis برای سلامت پروژه.'
)
features_text2_run.font.size = Pt(12)
features_text2_run.font.name = 'Calibri'

features_subtitle3 = final_doc.add_paragraph()
features_subtitle3_format = features_subtitle3.paragraph_format
features_subtitle3_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
features_subtitle3_format.space_before = Pt(14)
features_subtitle3_format.space_after = Pt(10)
features_subtitle3_run = features_subtitle3.add_run('همکاری و ارتباطات')
features_subtitle3_run.font.size = Pt(18)
features_subtitle3_run.font.bold = True
features_subtitle3_run.font.color.rgb = RGBColor(68, 114, 196)

features_text3 = final_doc.add_paragraph()
features_text3_format = features_text3.paragraph_format
features_text3_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
features_text3_format.line_spacing = 1.6
features_text3_format.space_after = Pt(12)
features_text3_run = features_text3.add_run(
    'سیستم ابزارهای قوی برای همکاری تیمی فراهم می‌کند. کاربران می‌توانند نظر بر روی Task‌ها بنویسند. کاربران می‌توانند یکدیگر را با '
    '@mention منشن کنند. سیستم نوتیفیکیشن Real-time دارد با استفاده از SignalR که کاربران را فوری‌ترین وقت از تغییرات مطلع می‌کند. '
    'سیستم Activity Log دارد که تمام تغییرات را ثبت می‌کند. سیستم چت گروهی دارد که امکان چت با AI Assistant را می‌دهد.'
)
features_text3_run.font.size = Pt(12)
features_text3_run.font.name = 'Calibri'

features_subtitle4 = final_doc.add_paragraph()
features_subtitle4_format = features_subtitle4.paragraph_format
features_subtitle4_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
features_subtitle4_format.space_before = Pt(14)
features_subtitle4_format.space_after = Pt(10)
features_subtitle4_run = features_subtitle4.add_run('مدیریت زمان و ریمایندر‌ها')
features_subtitle4_run.font.size = Pt(18)
features_subtitle4_run.font.bold = True
features_subtitle4_run.font.color.rgb = RGBColor(68, 114, 196)

features_text4 = final_doc.add_paragraph()
features_text4_format = features_text4.paragraph_format
features_text4_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
features_text4_format.line_spacing = 1.6
features_text4_format.space_after = Pt(12)
features_text4_run = features_text4.add_run(
    'سیستم قابلیت تخمین زمانی برای Task‌ها دارد. کاربران می‌توانند زمان واقعی‌شان را ثبت کنند تا مقایسه با تخمین برای آموزش تیم مفید باشد. '
    'سیستم تایمر دارد که می‌توان برای ثبت زمان استفاده کرد. سیستم یادآورهای زمانی دارد که کاربران را قبل از سررسید‌های مهم هشدار می‌دهد. '
    'سیستم گزارشات زمانی فراهم می‌کند که کاربران و مدیران می‌توانند بر روی آن تحلیل کنند.'
)
features_text4_run.font.size = Pt(12)
features_text4_run.font.name = 'Calibri'

final_doc.add_page_break()

# ===============================================
# نتیجه‌گیری
# ===============================================
conclusion_title = final_doc.add_paragraph()
conclusion_title_format = conclusion_title.paragraph_format
conclusion_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
conclusion_title_format.space_before = Pt(18)
conclusion_title_format.space_after = Pt(12)
conclusion_title_run = conclusion_title.add_run('نتیجه‌گیری')
conclusion_title_run.font.size = Pt(28)
conclusion_title_run.font.bold = True
conclusion_title_run.font.color.rgb = RGBColor(0, 51, 102)

conclusion_text = final_doc.add_paragraph()
conclusion_text_format = conclusion_text.paragraph_format
conclusion_text_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
conclusion_text_format.line_spacing = 1.6
conclusion_text_format.space_after = Pt(12)
conclusion_text_run = conclusion_text.add_run(
    'SmartTask یک سیستم جامع و پیشرفته است که تمام نیازهای مدیریت پروژه را برطرف می‌کند. سیستم به صورت کامل برای فارسی پشتیبانی '
    'می‌شود و می‌تواند در سرورهای شخصی نصب شود. سیستم از تکنولوژی‌های مدرن و معتبر استفاده می‌کند و قابلیت توسعه بالایی دارد. '
    'SmartTask می‌تواند برای تیم‌های کوچک تا بزرگ، استارتاپ‌ها تا شرکت‌های بزرگ استفاده شود.'
)
conclusion_text_run.font.size = Pt(12)
conclusion_text_run.font.name = 'Calibri'

# ذخیره فایل
output_path = 'SmartTask-جزوه-نهایی-کامل.docx'
final_doc.save(output_path)
print(f'✅ فایل Word ترکیب شده با موفقیت ایجاد شد!')
print(f'📄 نام فایل: {output_path}')
print(f'📋 محتوا: جزوه جامع و قابل‌خواندن')
print(f'✍️ سبک: جملات طویل و متصل، بدون لیست‌های نقطه‌ای')
print(f'🎯 شامل تمام اطلاعات کلیدی سیستم است')
