from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# خواندن فایل‌های موجود
doc1 = Document('SmartTask-Documentation.docx')
doc2 = Document('SmartTask-جزوه-کامل.docx')

# ایجاد document جدید
doc_final = Document()

# تنظیمات صفحه
section = doc_final.sections[0]
section.right_to_left = True
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# =====================================================
# صفحه عنوان
# =====================================================
title = doc_final.add_paragraph()
title_format = title.paragraph_format
title_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_format.space_before = Pt(60)
title_format.space_after = Pt(20)
title_run = title.add_run("راهنمای جامع سیستم SmartTask")
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc_final.add_paragraph()
subtitle_format = subtitle.paragraph_format
subtitle_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle_format.space_after = Pt(8)
subtitle_run = subtitle.add_run("سیستم مدیریت وظایف هوشمند مبتنی بر Agile/Scrum")
subtitle_run.font.size = Pt(16)
subtitle_run.font.bold = True

desc = doc_final.add_paragraph()
desc_format = desc.paragraph_format
desc_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
desc_run = desc.add_run("نسخه 1.0 - مستندات کامل، جامع و عملی")
desc_run.font.size = Pt(12)

doc_final.add_page_break()

# =====================================================
# معرفی پروژه
# =====================================================
h1 = doc_final.add_heading("درباره SmartTask", level=1)
h1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

intro = doc_final.add_paragraph(
    "SmartTask یک سیستم مدیریت وظایف و پروژه جامع است که برای تیم‌های توسعه نرم‌افزار طراحی شده است. "
    "این سیستم بر پایه معماری ASP.NET Core 8.0 ساخته شده و تمام فرآیندهای مدیریت پروژه را از برنامه‌ریزی "
    "اولیه تا تحویل نهایی پوشش می‌دهد. SmartTask یک راهکار یکپارچه، قدرتمند و کاربرپسند است که تیم‌ها "
    "را در برنامه‌ریزی و اجرای پروژه‌های خود یاری می‌دهد."
)
intro.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
intro.paragraph_format.line_spacing = 1.6
intro.paragraph_format.space_after = Pt(12)

features_intro = doc_final.add_paragraph(
    "یکی از ویژگی‌های متمایز SmartTask پشتیبانی کامل از زبان فارسی است. تمام قسمت‌های سیستم از رابط "
    "کاربری تا گزارش‌ها و تقویم‌ها به فارسی پشتیبانی می‌شوند. علاوه بر این، SmartTask می‌تواند به صورت "
    "Self-hosted در سرورهای شخصی نصب شود که این امکان کنترل کامل بر داده‌ها و عدم وابستگی به ابرهای "
    "خارجی را فراهم می‌کند. سیستم هوش مصنوعی داخلی‌اش (بدون وابستگی به API‌های خارجی) توانایی تجزیه "
    "خودکار وظایف و تحلیل‌های هوشمند را ارائه می‌دهد."
)
features_intro.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
features_intro.paragraph_format.line_spacing = 1.6
features_intro.paragraph_format.space_after = Pt(12)

doc_final.add_paragraph()

# =====================================================
# چرا SmartTask
# =====================================================
h2 = doc_final.add_heading("چرا SmartTask انتخاب کنیم", level=2)
h2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

why = doc_final.add_paragraph(
    "در مقایسه با ابزارهای مشابه مثل Jira، Trello و Asana، SmartTask مزایایی منحصربه‌فرد دارد. "
    "اول اینکه، SmartTask محلی است و می‌تواند در سرورهای خصوصی نصب شود، که بدان معنی است کنترل کامل "
    "داده‌ها در دست شما است و هیچ هزینه اشتراک ماهانه وجود ندارد. دوم اینکه، کد منبع سیستم در دسترس است "
    "و می‌توان آن را برای نیازهای خاص سازمان سفارشی‌سازی کرد. سوم اینکه، SmartTask پشتیبانی کامل از فارسی "
    "دارد که برای تیم‌های ایرانی بسیار مهم است. چهارم اینکه، سیستم هوش مصنوعی داخلی دارد و نیازی به API‌های "
    "خارجی نیست. پنجم اینکه، پس از نصب، می‌تواند در شبکه محلی بدون اینترنت استفاده شود."
)
why.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
why.paragraph_format.line_spacing = 1.6
why.paragraph_format.space_after = Pt(12)

doc_final.add_paragraph()

# =====================================================
# Agile و Scrum
# =====================================================
h2 = doc_final.add_heading("فلسفه Agile و Scrum", level=2)
h2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

agile_intro = doc_final.add_paragraph(
    "Agile یک روش‌شناسی نوین توسعه نرم‌افزار است که به جای برنامه‌ریزی کامل در ابتدا، بر توسعه تکراری "
    "و مرتب‌کار تمرکز دارد. در روش Agile، تیم‌ها به صورت مستمر بازخورد دریافت می‌کنند و بر اساس آن "
    "تغییرات لازم را اعمال می‌کنند. این روش‌شناسی بر انعطاف‌پذیری، همکاری تیمی و بازخورد مشتری تاکید دارد."
)
agile_intro.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
agile_intro.paragraph_format.line_spacing = 1.6
agile_intro.paragraph_format.space_after = Pt(12)

scrum_intro = doc_final.add_paragraph(
    "Scrum یک فریم‌ورک عملیاتی درون Agile است که روند توسعه را به بازه‌های زمانی کوتاه مدت (Sprint) "
    "تقسیم می‌کند. در هر Sprint، که معمولاً ۱ تا ۴ هفته است، تیم مجموعه‌ای از User Story‌ها را انتخاب "
    "کرده و در تلاش برای تکمیل آن‌ها است. Scrum چندین مراسم مهم دارد: Sprint Planning که در آن تیم "
    "نیازمندی‌ها را برای اسپرینت انتخاب می‌کند، Daily Standup که جلسات روزانه ۱۵ دقیقه‌ای است، Sprint Review "
    "که نتایج نمایش داده می‌شوند، و Sprint Retrospective که فرآیند بررسی می‌شود. SmartTask تمامی این "
    "فرآیندها را به صورت دیجیتال پیاده‌سازی کرده است."
)
scrum_intro.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
scrum_intro.paragraph_format.line_spacing = 1.6
scrum_intro.paragraph_format.space_after = Pt(12)

doc_final.add_paragraph()

# =====================================================
# معماری
# =====================================================
h2 = doc_final.add_heading("معماری و ساختار سیستم", level=2)
h2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

arch = doc_final.add_paragraph(
    "SmartTask بر اساس معماری N-Tier (چند لایه) طراحی شده است. این معماری مسئولیت‌ها را به صورت منطقی "
    "جدا می‌کند و هر لایه یک نقش مشخص دارد. لایه اول، لایه Presentation است که Controller‌ها، View‌ها و "
    "ViewModel‌ها را شامل می‌شود و مسئولیت دریافت درخواست‌های کاربر و نمایش نتایج را دارد. لایه دوم، لایه Service است "
    "که منطق کسب‌وکار را پیاده‌سازی می‌کند. لایه سوم، لایه Infrastructure است که Repository‌ها و UnitOfWork را شامل می‌شود "
    "و دسترسی به داده‌ها را مدیریت می‌کند. لایه چهارم، لایه Data است که Entity‌ها و Configuration‌ها را دارد. و بالاخره "
    "لایه Database است که پایگاه داده SQL Server است."
)
arch.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
arch.paragraph_format.line_spacing = 1.6
arch.paragraph_format.space_after = Pt(12)

patterns = doc_final.add_paragraph(
    "SmartTask از چندین الگوی طراحی معروف استفاده می‌کند تا کد را بهتر سازمان‌دهی کند. الگوی Repository برای "
    "انتزاع لایه دسترسی به داده استفاده می‌شود و تست‌پذیری بیشتری را فراهم می‌کند. الگوی Unit of Work برای مدیریت "
    "Transaction‌ها و اطمینان از یکپارچگی داده استفاده می‌شود. Dependency Injection برای کاهش وابستگی‌ها و افزایش انعطاف‌پذیری "
    "استفاده می‌شود. Service Layer Pattern برای جداسازی منطق کسب‌وکار از Controller استفاده می‌شود. Soft Delete Pattern برای حذف "
    "منطقی به جای حذف فیزیکی رکوردها استفاده می‌شود. و Hub Pattern برای SignalR Hubs استفاده می‌شود."
)
patterns.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
patterns.paragraph_format.line_spacing = 1.6
patterns.paragraph_format.space_after = Pt(12)

doc_final.add_page_break()

# =====================================================
# تکنولوژی‌ها
# =====================================================
h2 = doc_final.add_heading("فناوری‌های استفاده شده", level=2)
h2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

tech_intro = doc_final.add_paragraph(
    "SmartTask بر پایه فناوری‌های مدرن و معتبری ساخته شده است. فریم‌ورک اصلی ASP.NET Core 8.0 است که "
    "توسط مایکروسافت توسعه یافته و کاملاً متن‌باز است. این فریم‌ورک دارای معماری MVC، Cross-platform "
    "(Windows, Linux, macOS)، عملکرد بالا و مناسبیت برای پروژه‌های بزرگ است."
)
tech_intro.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
tech_intro.paragraph_format.line_spacing = 1.6
tech_intro.paragraph_format.space_after = Pt(12)

db = doc_final.add_paragraph(
    "برای پایگاه داده، SQL Server استفاده می‌شود که یک سیستم مدیریت پایگاه داده رابطه‌ای قدرتمند از "
    "مایکروسافت است. Entity Framework Core 8.0 ORM (Object-Relational Mapping) استفاده شده است که امکان کار "
    "با پایگاه داده را از طریق کلاس‌های .NET فراهم می‌کند. برای احراز هویت، ASP.NET Core Identity استفاده می‌شود که "
    "سیستمی جامع برای مدیریت کاربران، پسورد‌ها و نقش‌ها است."
)
db.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
db.paragraph_format.line_spacing = 1.6
db.paragraph_format.space_after = Pt(12)

realtime = doc_final.add_paragraph(
    "برای ارتباطات Real-time، SignalR استفاده می‌شود که امکان پیام‌های دوطرفه بدون تاخیر را فراهم می‌کند. "
    "ClosedXML برای تولید فایل‌های Excel، QuestPDF برای تولید فایل‌های PDF استفاده می‌شود. برای تست، xUnit و Moq "
    "استفاده می‌شود. در Frontend، Bootstrap 5 برای طراحی Responsive UI، jQuery برای کار با DOM و AJAX، و Chart.js برای نمودارها استفاده می‌شود."
)
realtime.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
realtime.paragraph_format.line_spacing = 1.6
realtime.paragraph_format.space_after = Pt(12)

ai = doc_final.add_paragraph(
    "برای هوش مصنوعی، LM Studio استفاده می‌شود که یک مدل AI محلی (Google Gemma) را اجرا می‌کند. این سیستم "
    "بدون وابستگی به API‌های خارجی عمل می‌کند و حریم خصوصی کاملی را تضمین می‌کند."
)
ai.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
ai.paragraph_format.line_spacing = 1.6
ai.paragraph_format.space_after = Pt(12)

doc_final.add_paragraph()

# =====================================================
# موجودیت‌های کلیدی
# =====================================================
h2 = doc_final.add_heading("اجزای کلیدی سیستم", level=2)
h2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

h3 = doc_final.add_heading("Workspace (فضای کاری)", level=3)
h3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

ws = doc_final.add_paragraph(
    "Workspace بالاترین سطح سازمان‌دهی در سیستم است. هر Workspace متعلق به یک Owner است و می‌تواند شامل "
    "چندین Project، Team و Member باشد. یک شرکت می‌تواند یک Workspace داشته باشد یا چندین Workspace برای "
    "بخش‌های مختلف. هر Workspace دارای نام، توضیحات، لوگو و رنگ سفارشی است. در Workspace می‌توان مخاطبان "
    "و نقش‌ها را مدیریت کرد با نقش‌هایی مثل Owner (مالک)، Admin (مدیر)، ProjectManager و Viewer."
)
ws.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
ws.paragraph_format.line_spacing = 1.6
ws.paragraph_format.space_after = Pt(12)

h3 = doc_final.add_heading("Project (پروژه)", level=3)
h3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

proj = doc_final.add_paragraph(
    "Project پروژه‌های نرم‌افزاری را نمایندگی می‌کند و محل اصلی برنامه‌ریزی و اجرای کار است. هر Project در داخل "
    "یک Workspace قرار دارد و دارای نام، کلید یکتا، توضیحات، رنگ و آیکون است. هر Project دارای تاریخ شروع، سررسید و "
    "تاریخ پایان است. Project دارای وضعیت مختلفی است: Planning، InProgress، OnHold، Completed و Cancelled. هر Project "
    "دارای اولویت است: Low، Medium، High و Critical."
)
proj.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
proj.paragraph_format.line_spacing = 1.6
proj.paragraph_format.space_after = Pt(12)

h3 = doc_final.add_heading("Sprint (اسپرینت)", level=3)
h3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

sprint = doc_final.add_paragraph(
    "Sprint بازه زمانی مشخصی است که تیم درآن مجموعه‌ای از User Story‌ها را تکمیل کرتا تلاش می‌کند. معمولاً ۱ تا ۴ هفته طول می‌کشد. "
    "هر Sprint دارای نام، هدف، تاریخ شروع و پایان است. هر Sprint دارای ظرفیت (Capacity) است که به ساعت اندازه‌گیری می‌شود. "
    "Sprint دارای وضعیت مختلفی است: Planning، Active، Completed و Cancelled. تنها یک Sprint می‌تواند در هر Project فعال باشد."
)
sprint.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
sprint.paragraph_format.line_spacing = 1.6
sprint.paragraph_format.space_after = Pt(12)

h3 = doc_final.add_heading("UserStory (داستان کاربری)", level=3)
h3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

story = doc_final.add_paragraph(
    "UserStory نیازمندی یا ویژگی‌ای است که از دیدگاه کاربر نوشته شده است. UserStory واحد کاری در Scrum است و معمولاً "
    "شامل چندین Task است. هر UserStory دارای عنوان، توضیحات و معیارهای پذیرش است که شرایطی هستند که Story باید آن‌ها را برآورده کند. "
    "هر UserStory دارای Story Point است که واحد تخمین پیچیدگی است (معمولاً بر اساس سری فیبوناچی). UserStory دارای وضعیت مختلفی است: "
    "New، InProgress، Testing و Done. UserStory می‌تواند در Backlog یا در Sprint قرار داشته باشد."
)
story.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
story.paragraph_format.line_spacing = 1.6
story.paragraph_format.space_after = Pt(12)

h3 = doc_final.add_heading("TaskItem (وظیفه)", level=3)
h3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

task = doc_final.add_paragraph(
    "TaskItem کوچک‌ترین واحد کاری قابل تخصیص است که باید به یک UserStory متصل باشد. پیاده‌سازی یک UserStory معمولاً شامل "
    "چندین Task است. هر Task دارای عنوان، توضیحات و نوع است که می‌تواند Task (کار عادی)، Bug (اشکال)، Feature (قابلیت جدید) "
    "یا Improvement (بهبود) باشد. Task دارای وضعیت مختلفی است: ToDo، InProgress، InReview، Done و Cancelled. هر Task دارای تخمین زمانی "
    "به ساعت است و می‌تواند تاریخ شروع، سررسید و تکمیل داشته باشد. Task می‌تواند به چندین نفر تخصیص داده شود و می‌تواند SubTask، "
    "Comment، Attachment، Checklist، Label و وابستگی داشته باشد."
)
task.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
task.paragraph_format.line_spacing = 1.6
task.paragraph_format.space_after = Pt(12)

doc_final.add_page_break()

# =====================================================
# قابلیت‌ها
# =====================================================
h2 = doc_final.add_heading("قابلیت‌های اصلی سیستم", level=2)
h2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

cap1_h3 = doc_final.add_heading("مدیریت کاربران و دسترسی", level=3)
cap1_h3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

cap1 = doc_final.add_paragraph(
    "سیستم SmartTask یک سیستم احراز هویت و مجوز جامع دارد. کاربران می‌توانند با نام کاربری و رمز عبور ثبت‌نام و وارد سیستم شوند. "
    "سیستم پشتیبانی از احراز هویت دو مرحله‌ای دارد برای امنیت بیشتر. امکان ورود با Google نیز وجود دارد (فعلاً غیرفعال). "
    "نقش‌های سیستمی شامل Admin و User است. نقش‌های Workspace شامل Owner، Admin، ProjectManager، Developer، Tester و Viewer است. "
    "نقش‌های Project شامل Manager، Developer، Tester و Viewer است."
)
cap1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
cap1.paragraph_format.line_spacing = 1.6
cap1.paragraph_format.space_after = Pt(12)

cap2_h3 = doc_final.add_heading("داشبوردها و گزارشات", level=3)
cap2_h3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

cap2 = doc_final.add_paragraph(
    "سیستم داشبوردهای متنوعی برای مختلف سطح‌های کاربران فراهم می‌کند. Dashboard کاربر نمای کلی وظایف و فعالیت‌های شخصی را نمایش می‌دهد. "
    "Dashboard Workspace آمار پروژه‌ها و اعضا را نشان می‌دهد. Dashboard Project پیشرفت، نمودارها و KPI‌ها را نمایش می‌دهد. "
    "Dashboard Admin آمار کل سیستم را نشان می‌دهد. گزارشات تحلیلی پروژه، Sprint و Workspace فراهم می‌شود. نمودار Burndown و Velocity "
    "برای ارزیابی پیشرفت استفاده می‌شود. تحلیل سلامت پروژه (Health)، ریسک تاخیر (Delay Risk) و تحلیل بار کاری (Workload Analysis) "
    "ارائه می‌شود."
)
cap2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
cap2.paragraph_format.line_spacing = 1.6
cap2.paragraph_format.space_after = Pt(12)

cap3_h3 = doc_final.add_heading("تابلوها و نماها", level=3)
cap3_h3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

cap3 = doc_final.add_paragraph(
    "سیستم تابلوهای مختلفی برای مدیریت وظایف فراهم می‌کند. Task Board نمای Kanban با Drag & Drop را فراهم می‌کند. "
    "Backlog برای مدیریت UserStory‌های بدون Sprint است. Sprint Planning برای جابجایی Story بین Backlog و Sprint استفاده می‌شود. "
    "فیلتر‌های پیشرفته برای وضعیت، اولویت، تخصیص و برچسب در دسترس هستند. نمای Timeline (Gantt-like) و نمای تقویم نیز موجود است."
)
cap3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
cap3.paragraph_format.line_spacing = 1.6
cap3.paragraph_format.space_after = Pt(12)

cap4_h3 = doc_final.add_heading("همکاری و ارتباطات", level=3)
cap4_h3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

cap4 = doc_final.add_paragraph(
    "SmartTask ابزارهای قدرتمندی برای همکاری تیمی فراهم می‌کند. کاربران می‌توانند نظرات (Comments) بر روی Task اضافه کنند. "
    "منشن کردن اعضا (@username) برای جلب توجه آن‌ها کار می‌کند. نوتیفیکیشن Real-time با SignalR فراهم می‌شود. Activity Log تمام تغییرات را ثبت می‌کند. "
    "چت گروهی با AI Assistant (LM Studio) موجود است. ارسال دعوتنامه به Workspace و اضافه کردن اعضا جدید ممکن است."
)
cap4.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
cap4.paragraph_format.line_spacing = 1.6
cap4.paragraph_format.space_after = Pt(12)

cap5_h3 = doc_final.add_heading("مدیریت زمان", level=3)
cap5_h3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

cap5 = doc_final.add_paragraph(
    "سیستم ابزارهایی برای مدیریت زمان فراهم می‌کند. کاربران می‌توانند زمان کار (TimeLog) بر روی Task ثبت کنند. "
    "تایمر شروع/توقف برای پیگیری زمان واقعی موجود است. یادآورهای زمانی (Reminder) برای مهلت‌های مهم قابل تنظیم هستند. "
    "گزارش زمان صرف شده به ازای کاربر، پروژه و Task فراهم می‌شود. مقایسه زمان تخمینی و واقعی کمک می‌کند تا تخمین‌های بهتری انجام شود."
)
cap5.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
cap5.paragraph_format.line_spacing = 1.6
cap5.paragraph_format.space_after = Pt(12)

doc_final.add_paragraph()

# =====================================================
# ویژگی‌های پیشرفته
# =====================================================
h2 = doc_final.add_heading("ویژگی‌های پیشرفته", level=2)
h2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

adv = doc_final.add_paragraph(
    "SmartTask چندین ویژگی پیشرفته دارد که مدیریت پروژه‌ها را بیشتر تسهیل می‌کند. وابستگی‌های Task (Dependencies) امکان "
    "تعریف روابط بین Task‌ها را می‌دهد مثل Blocked By و Blocks. مبادله Task بین Project‌ها (Task Trade) برای انتقال وظایف "
    "بین پروژه‌ها کار می‌کند. Task Breakdown با کمک AI توانایی تقسیم خودکار Task‌های بزرگ را فراهم می‌کند. تحلیل اولویت هوشمند "
    "(Priority Engine) کمک می‌کند تا Task‌ها بر اساس اهمیت و فوریت دسته‌بندی شوند."
)
adv.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
adv.paragraph_format.line_spacing = 1.6
adv.paragraph_format.space_after = Pt(12)

adv2 = doc_final.add_paragraph(
    "Offroad Task برای وظایف خارج از پروژه استفاده می‌شود. برچسب‌گذاری رنگی (Labels) برای دسته‌بندی Task‌ها کار می‌کند. "
    "Checklist در Task برای نقاط کنترل کوچک‌تر استفاده می‌شود. SubTask برای تقسیم Task‌ها به واحد‌های کوچک‌تر کار می‌کند. "
    "ضمیمه فایل (Attachments) اسناد و فایل‌های مرتبط را به Task اضافه می‌کند."
)
adv2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
adv2.paragraph_format.line_spacing = 1.6
adv2.paragraph_format.space_after = Pt(12)

doc_final.add_page_break()

# =====================================================
# سرویس‌ها
# =====================================================
h2 = doc_final.add_heading("سرویس‌های سیستم", level=2)
h2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

services = doc_final.add_paragraph(
    "SmartTask از ۴۵+ سرویس متخصص استفاده می‌کند. WorkspaceService برای مدیریت فضاهای کاری، ProjectService برای پروژه‌ها، "
    "SprintService برای Sprint‌ها، UserStoryService برای داستان‌های کاربری، TaskService برای وظایف استفاده می‌شود. "
    "NotificationService سیستم نوتیفیکیشن را مدیریت می‌کند. ActivityLogService تمام فعالیت‌ها را ثبت می‌کند. "
    "TaskAssignmentService تخصیص وظایف را مدیریت می‌کند. TimeLogService ثبت زمان کار، CommentService مدیریت نظرات، "
    "AttachmentService فایل‌های ضمیمه و LabelService برچسب‌ها را مدیریت می‌کند."
)
services.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
services.paragraph_format.line_spacing = 1.6
services.paragraph_format.space_after = Pt(12)

services2 = doc_final.add_paragraph(
    "ChecklistService چک‌لیست‌ها، SubTaskService زیروظایف، ReminderService یادآورها، TaskDependencyService وابستگی‌ها "
    "را مدیریت می‌کند. WorkloadAnalysisService تحلیل بار کاری، DelayRiskService تحلیل ریسک تاخیر، ProjectHealthService تحلیل "
    "سلامت پروژه انجام می‌دهد. PriorityEngineService موتور اولویت‌بندی هوشمند، ReportExportService خروجی گزارشات، ChatService چت "
    "با AI و WebpushrService ارسال Push Notification را مدیریت می‌کند."
)
services2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
services2.paragraph_format.line_spacing = 1.6
services2.paragraph_format.space_after = Pt(12)

doc_final.add_paragraph()

# =====================================================
# پایان
# =====================================================
h2 = doc_final.add_heading("نتیجه‌گیری", level=2)
h2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

conclusion = doc_final.add_paragraph(
    "SmartTask یک سیستم جامع و قدرتمند برای مدیریت پروژه‌ها است که تمام نیازهای تیم‌های توسعه نرم‌افزار را پوشش می‌دهد. "
    "معماری آن بر اساس بهترین عملکردهای صنعت است و استفاده از آن برای تیم‌های کوچک تا بزرگ مناسب است. "
    "پشتیبانی کامل از فارسی، قابلیت Self-hosting و هوش مصنوعی داخلی آن را از رقبا متمایز می‌کند. "
    "SmartTask نه تنها یک ابزار است، بلکه یک راهکار جامع برای بهبود بهره‌وری تیمی و مدیریت پروژه‌های موفق است."
)
conclusion.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
conclusion.paragraph_format.line_spacing = 1.6
conclusion.paragraph_format.space_after = Pt(12)

# ذخیره فایل
output_path = 'SmartTask-جزوه-نهایی.docx'
doc_final.save(output_path)
print(f'✅ فایل جزوه نهایی با موفقیت ایجاد شد: {output_path}')
print(f'📄 فایل جمله‌محور و بدون پیچیدگی متنی است')
