"""
تبدیل جزوه‌های مارک‌داون فارسی به HTML / PDF / DOCX با پشتیبانی کامل RTL
PDF از طریق Microsoft Edge در حالت headless ساخته می‌شود.
"""
import re
import shutil
import subprocess
import sys
from pathlib import Path

# کنسول ویندوز پیش‌فرض cp1256 است و فارسی را چاپ نمی‌کند
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

BASE_DIR = Path(r"E:\taskManager\SmartTaskManagementSystem\SmartTask Learn")

# پروفایل موقت مرورگر؛ بدون آن حالت headless گاهی خروجی نمی‌دهد
PROFILE_DIR = Path(r"C:\Users\User\AppData\Local\Temp\smarttask_pdf_profile")

EDGE_CANDIDATES = [
    r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
    r"C:\Program Files\Google\Chrome\Application\chrome.exe",
    r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
]

CSS = """
@page { size: A4; margin: 1.8cm 1.6cm; }

body {
    font-family: 'Vazirmatn', 'Segoe UI', 'Tahoma', sans-serif;
    direction: rtl; text-align: right;
    line-height: 1.9; color: #1f2937;
    font-size: 11pt; margin: 0; padding: 0;
}

h1 { font-size: 20pt; color: #1e3a8a; border-bottom: 3px solid #3b82f6;
     padding-bottom: .3em; margin: 1.2em 0 .8em; page-break-after: avoid; }
h2 { font-size: 15pt; color: #1e40af; border-bottom: 1px solid #cbd5e1;
     padding-bottom: .25em; margin: 1.4em 0 .6em; page-break-after: avoid; }
h3 { font-size: 12.5pt; color: #334155; margin: 1.1em 0 .5em; page-break-after: avoid; }
h4 { font-size: 11.5pt; color: #475569; margin: 1em 0 .4em; page-break-after: avoid; }

p { margin: .6em 0; text-align: justify; }

code {
    background: #f1f5f9; color: #0f172a; padding: 1px 5px;
    border-radius: 3px; font-family: 'Consolas', monospace;
    font-size: 9.5pt; direction: ltr; unicode-bidi: embed;
}

pre {
    background: #0f172a; color: #e2e8f0; padding: 12px 14px;
    border-radius: 6px; overflow-x: auto; direction: ltr; text-align: left;
    font-family: 'Consolas', monospace; font-size: 9pt; line-height: 1.55;
    page-break-inside: avoid; margin: .8em 0;
}
pre code { background: transparent; color: inherit; padding: 0; font-size: inherit; }

table {
    width: 100%; border-collapse: collapse; margin: .9em 0;
    font-size: 10pt; page-break-inside: avoid; direction: rtl;
}
th, td { border: 1px solid #cbd5e1; padding: 7px 9px; text-align: right; vertical-align: top; }
th { background: #3b82f6; color: #fff; font-weight: 700; }
tr:nth-child(even) td { background: #f8fafc; }

ul, ol { margin: .6em 0; padding-right: 1.6em; padding-left: 0; }
li { margin: .3em 0; }

blockquote {
    border-right: 4px solid #3b82f6; background: #f8fafc;
    padding: .6em 1em; margin: .9em 0; color: #475569;
}

hr { border: none; border-top: 1px solid #e2e8f0; margin: 1.6em 0; }
strong { color: #1e3a8a; font-weight: 700; }
a { color: #2563eb; text-decoration: none; }
"""


def render_inline(text: str) -> str:
    """پردازش نشانه‌گذاری درون‌خطی (کد، بولد، ایتالیک، لینک)."""
    placeholders: list[str] = []

    def stash(match: re.Match) -> str:
        placeholders.append(match.group(1))
        return f"\x00{len(placeholders) - 1}\x00"

    # کد درون‌خطی اول کنار گذاشته می‌شود تا محتوایش پردازش نشود
    text = re.sub(r"`([^`]+)`", stash, text)

    text = (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)", r"<em>\1</em>", text)
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r'<a href="\2">\1</a>', text)

    for i, raw in enumerate(placeholders):
        safe = raw.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        text = text.replace(f"\x00{i}\x00", f"<code>{safe}</code>")

    return text


def markdown_to_html(md_text: str) -> str:
    """تبدیل مارک‌داون به HTML — سطر به سطر تا جدول و کدبلاک درست بمانند."""
    out: list[str] = []
    lines = md_text.split("\n")
    i = 0
    in_code = False
    list_stack: list[str] = []

    def close_lists() -> None:
        while list_stack:
            out.append(f"</{list_stack.pop()}>")

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # کدبلاک
        if stripped.startswith("```"):
            if in_code:
                out.append("</code></pre>")
                in_code = False
            else:
                close_lists()
                out.append("<pre><code>")
                in_code = True
            i += 1
            continue

        if in_code:
            out.append(line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
            i += 1
            continue

        # جدول: سطر فعلی و سطر جداکننده بعدی
        if (
            stripped.startswith("|")
            and i + 1 < len(lines)
            and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip())
        ):
            close_lists()
            headers = [c.strip() for c in stripped.strip("|").split("|")]
            out.append("<table><thead><tr>")
            out.extend(f"<th>{render_inline(h)}</th>" for h in headers)
            out.append("</tr></thead><tbody>")
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                out.append("<tr>")
                out.extend(f"<td>{render_inline(c)}</td>" for c in cells)
                out.append("</tr>")
                i += 1
            out.append("</tbody></table>")
            continue

        if not stripped:
            close_lists()
            i += 1
            continue

        # سرتیترها
        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            close_lists()
            level = len(heading.group(1))
            out.append(f"<h{level}>{render_inline(heading.group(2))}</h{level}>")
            i += 1
            continue

        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", stripped):
            close_lists()
            out.append("<hr>")
            i += 1
            continue

        if stripped.startswith(">"):
            close_lists()
            out.append(f"<blockquote>{render_inline(stripped.lstrip('> '))}</blockquote>")
            i += 1
            continue

        # فهرست نامرتب
        bullet = re.match(r"^[-*+]\s+(.*)$", stripped)
        if bullet:
            if "ul" not in list_stack:
                close_lists()
                out.append("<ul>")
                list_stack.append("ul")
            out.append(f"<li>{render_inline(bullet.group(1))}</li>")
            i += 1
            continue

        # فهرست مرتب
        numbered = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if numbered:
            if "ol" not in list_stack:
                close_lists()
                out.append("<ol>")
                list_stack.append("ol")
            out.append(f"<li>{render_inline(numbered.group(1))}</li>")
            i += 1
            continue

        close_lists()
        out.append(f"<p>{render_inline(stripped)}</p>")
        i += 1

    if in_code:
        out.append("</code></pre>")
    close_lists()
    return "\n".join(out)


def build_html(md_path: Path, html_path: Path, title: str) -> None:
    md_text = md_path.read_text(encoding="utf-8")
    body = markdown_to_html(md_text)
    html_path.write_text(
        "<!DOCTYPE html>\n"
        '<html lang="fa" dir="rtl">\n<head>\n<meta charset="UTF-8">\n'
        f"<title>{title}</title>\n<style>{CSS}</style>\n</head>\n"
        f"<body>\n{body}\n</body>\n</html>\n",
        encoding="utf-8",
    )
    print(f"  [ok] HTML  -> {html_path.name}")


def find_browser() -> str | None:
    for candidate in EDGE_CANDIDATES:
        if Path(candidate).exists():
            return candidate
    return None


def build_pdf(html_path: Path, pdf_path: Path, browser: str | None) -> bool:
    if not browser:
        print("  [skip] PDF: هیچ مرورگری برای رندر پیدا نشد")
        return False

    if pdf_path.exists():
        pdf_path.unlink()

    # نکته: حالت قدیمی --headless بی‌صدا هیچ فایلی نمی‌سازد؛
    # حالت new همراه پروفایل جدا و بودجه زمانی لازم است.
    subprocess.run(
        [
            browser,
            "--headless=new",
            "--disable-gpu",
            "--no-sandbox",
            f"--user-data-dir={PROFILE_DIR}",
            "--virtual-time-budget=10000",
            "--no-pdf-header-footer",
            f"--print-to-pdf={pdf_path}",
            html_path.as_uri(),
        ],
        capture_output=True,
        timeout=240,
    )

    if pdf_path.exists() and pdf_path.stat().st_size > 2000:
        size_kb = pdf_path.stat().st_size / 1024
        print(f"  [ok] PDF   -> {pdf_path.name} ({size_kb:.0f} KB)")
        return True

    print("  [fail] PDF ساخته نشد")
    return False


def build_docx(md_path: Path, docx_path: Path) -> bool:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        print("  [skip] DOCX: بسته python-docx نصب نیست")
        return False

    def set_rtl(paragraph) -> None:
        p = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        p.append(bidi)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def style_runs(paragraph, size=11, bold=False, color=None, mono=False) -> None:
        for run in paragraph.runs:
            run.font.name = "Consolas" if mono else "Vazirmatn"
            run.font.size = Pt(size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color
            rpr = run._element.get_or_add_rPr()
            rtl = OxmlElement("w:rtl")
            rtl.set(qn("w:val"), "0" if mono else "1")
            rpr.append(rtl)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    for attr in ("left_margin", "right_margin"):
        setattr(section, attr, Inches(0.7))

    lines = md_path.read_text(encoding="utf-8").split("\n")
    i = 0
    in_code = False
    code_buffer: list[str] = []

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                if code_buffer:
                    p = doc.add_paragraph("\n".join(code_buffer))
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    style_runs(p, size=9, mono=True)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(raw)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # جدول
        if (
            stripped.startswith("|")
            and i + 1 < len(lines)
            and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip())
        ):
            headers = [c.strip() for c in stripped.strip("|").split("|")]
            rows = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
                j += 1

            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Light Grid Accent 1"
            table.table_direction = WD_SECTION.NEW_PAGE if False else table.table_direction
            for idx, header in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.text = re.sub(r"\*\*|`", "", header)
                for p in cell.paragraphs:
                    set_rtl(p)
                    style_runs(p, size=10, bold=True)
            for row in rows:
                cells = table.add_row().cells
                for idx, value in enumerate(row[: len(headers)]):
                    cells[idx].text = re.sub(r"\*\*|`", "", value)
                    for p in cells[idx].paragraphs:
                        set_rtl(p)
                        style_runs(p, size=10)
            doc.add_paragraph()
            i = j
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = re.sub(r"\*\*|`", "", heading.group(2))
            p = doc.add_heading(text, level=min(level, 4))
            set_rtl(p)
            style_runs(p, size=max(19 - level * 2, 11), bold=True,
                       color=RGBColor(0x1E, 0x3A, 0x8A))
            i += 1
            continue

        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", stripped):
            i += 1
            continue

        clean = re.sub(r"\*\*|`|^[-*+]\s+|^\d+[.)]\s+|^>\s*", "", stripped)
        bullet = bool(re.match(r"^[-*+]\s+", stripped))
        numbered = bool(re.match(r"^\d+[.)]\s+", stripped))

        style = "List Bullet" if bullet else "List Number" if numbered else None
        p = doc.add_paragraph(clean, style=style) if style else doc.add_paragraph(clean)
        set_rtl(p)
        style_runs(p, size=11)
        i += 1

    doc.save(docx_path)
    size_kb = docx_path.stat().st_size / 1024
    print(f"  [ok] DOCX  -> {docx_path.name} ({size_kb:.0f} KB)")
    return True


DOCUMENTS = [
    ("جزوه 3 - سیستم گیمیفیکیشن", "جزوه ۳ — سیستم گیمیفیکیشن"),
    ("راهنمای بهره‌وری و شبیه‌سازی", "راهنمای بهره‌وری و شبیه‌سازی پروژه"),
]


def main() -> int:
    browser = find_browser()
    print(f"مرورگر رندر: {browser or 'پیدا نشد'}\n")

    failures = 0
    for stem, title in DOCUMENTS:
        md_path = BASE_DIR / f"{stem}.md"
        if not md_path.exists():
            print(f"[!] فایل مبدأ پیدا نشد: {md_path.name}")
            failures += 1
            continue

        print(f"» {stem}")
        html_path = BASE_DIR / f"{stem}.html"
        build_html(md_path, html_path, title)
        if not build_pdf(html_path, BASE_DIR / f"{stem}.pdf", browser):
            failures += 1
        if not build_docx(md_path, BASE_DIR / f"{stem}.docx"):
            failures += 1
        print()

    print("همه فایل‌ها ساخته شد." if not failures else f"{failures} مورد ناتمام ماند.")
    return 1 if failures else 0


if __name__ == "__main__":
    sys.exit(main())
